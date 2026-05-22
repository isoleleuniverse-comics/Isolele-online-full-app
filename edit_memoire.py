from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph


SRC = Path(r"C:\Users\ngese\Downloads\memore final.docx")
OUT = Path(r"C:\Users\ngese\isolele\refactor-isolele\memoire_corrige.docx")


def normalize(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").replace("’", "'").split())


def iter_paragraphs(doc: DocumentObject) -> list[Paragraph]:
    return list(doc.paragraphs)


def find_paragraph(doc: DocumentObject, needle: str, *, exact: bool = False) -> Paragraph:
    target = normalize(needle)
    for paragraph in iter_paragraphs(doc):
        current = normalize(paragraph.text)
        if (current == target) if exact else (target in current):
            return paragraph
    raise ValueError(f"Paragraph not found: {needle!r}")


def replace_text(paragraph: Paragraph, text: str) -> Paragraph:
    paragraph.text = text
    return paragraph


def block_xml(block: Paragraph | Table):
    return block._p if isinstance(block, Paragraph) else block._tbl


def block_parent(block: Paragraph | Table):
    return block._parent


def insert_paragraph_after(block: Paragraph | Table, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    block_xml(block).addnext(new_p)
    paragraph = Paragraph(new_p, block_parent(block))
    if text:
        paragraph.text = text
    if style is not None:
        paragraph.style = style
    return paragraph


def add_table_after(doc: DocumentObject, block: Paragraph | Table, rows: int, cols: int) -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    block_xml(block).addnext(table._tbl)
    return table


def remove_table(table: Table) -> None:
    table._tbl.getparent().remove(table._tbl)


def make_two_col_table(
    doc: DocumentObject,
    anchor: Paragraph | Table,
    style,
    title: str,
    rows: list[tuple[str, str]],
) -> Table:
    table = add_table_after(doc, anchor, rows=len(rows) + 1, cols=2)
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(1.75), Inches(4.85)]

    header = table.rows[0].cells
    header[0].text = "Cas d’utilisation"
    header[1].text = title

    for idx, (label, value) in enumerate(rows, start=1):
        cells = table.rows[idx].cells
        cells[0].text = label
        cells[1].text = value

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if cell.paragraphs:
                if idx == 0 and cell.paragraphs[0].runs:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    return table


def main() -> None:
    copyfile(SRC, OUT)
    doc = Document(OUT)

    # Acteurs et cadrage admin/MVP.
    replace_text(
        find_paragraph(doc, "Afin de maintenir le système"),
        "Afin de maintenir le système, l’administrateur est chargé de la gestion technique et fonctionnelle de la plateforme. Il supervise les comptes utilisateurs, peut intervenir sur leur accès en cas d’incident, consulte les statistiques globales d’usage et veille à la disponibilité de l’API REST ainsi qu’à l’intégrité des données. Dans le périmètre du MVP présenté au chapitre 3, cet acteur n’est pas encore matérialisé par une interface complète, mais il reste indispensable dans le modèle conceptuel du système.",
    )

    # Cas d'utilisation: remplacer le grand tableau par des fiches plus lisibles.
    caption = find_paragraph(doc, "Tableau 2 : Documentation des cas d’utilisation")
    original_table = doc.tables[1]
    original_style = original_table.style
    remove_table(original_table)

    student_intro = insert_paragraph_after(
        caption,
        "Les cas d’utilisation de l’étudiant sont présentés ci-dessous sous forme de fiches séparées afin de rendre chaque scénario plus lisible.",
        style=caption.style,
    )
    student_cases = [
        (
            "CU1 : S’authentifier",
            [
                ("Acteurs", "Étudiant, administrateur"),
                ("Description", "Permet à un acteur autorisé d’ouvrir une session sécurisée afin d’accéder à ses fonctionnalités et à ses données."),
                ("Préconditions", "Un compte existe et aucune session valide n’est active sur le poste concerné."),
                ("Postconditions", "Une session authentifiée est établie et les jetons JWT nécessaires sont générés."),
                ("Scénario nominal", "L’acteur saisit son adresse électronique et son mot de passe. Le système valide les identifiants, génère les jetons d’accès et redirige l’utilisateur vers l’interface correspondante."),
            ],
        ),
        (
            "CU2 : Gérer ses cours",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet à l’étudiant d’ajouter, de modifier, d’archiver ou de consulter ses unités d’enseignement."),
                ("Préconditions", "L’étudiant est authentifié."),
                ("Postconditions", "La liste des cours et leurs indicateurs associés sont mis à jour."),
                ("Scénario nominal", "L’étudiant ouvre la rubrique « Cours », ajoute un nouveau cours ou modifie un cours existant, puis le système enregistre les changements et actualise la liste affichée."),
            ],
        ),
        (
            "CU3 : Gérer les événements académiques",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet de créer, modifier ou supprimer des événements liés aux cours, aux devoirs ou aux examens."),
                ("Préconditions", "L’étudiant est authentifié et dispose d’au moins un cours auquel rattacher l’événement."),
                ("Postconditions", "Le calendrier de l’étudiant reflète les nouvelles échéances."),
                ("Scénario nominal", "L’étudiant ouvre le planning, choisit une date, crée un événement en précisant son type, son horaire et le cours associé, puis le système l’enregistre et l’affiche dans la vue calendrier."),
            ],
        ),
        (
            "CU4 : Gérer les tâches",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet de créer, classer, modifier et suivre l’avancement des tâches personnelles ou académiques."),
                ("Préconditions", "L’étudiant est authentifié."),
                ("Postconditions", "Le tableau des tâches et les statuts associés sont mis à jour."),
                ("Scénario nominal", "L’étudiant ajoute une tâche depuis le tableau Kanban ou la capture rapide, précise sa priorité, son échéance éventuelle et son lien avec un cours ou un événement, puis suit son passage entre les états « À faire », « En cours », « Terminées » ou « Annulées »."),
            ],
        ),
        (
            "CU5 : Enregistrer une note",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet d’ajouter une note obtenue à un travail, une interrogation, un TP ou un examen rattaché à un cours."),
                ("Préconditions", "L’étudiant est authentifié et a déjà créé au moins un cours."),
                ("Postconditions", "La note est enregistrée et les moyennes du cours concerné sont recalculées."),
                ("Scénario nominal", "L’étudiant ouvre la fiche d’un cours, choisit l’onglet des notes, saisit le score, le barème et le type de travail, puis le système sauvegarde la note et met à jour les indicateurs de performance."),
            ],
        ),
        (
            "CU6 : Consulter le risque académique",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet de visualiser, pour chaque cours, un niveau de risque calculé à partir des notes, des tâches en attente et de la proximité des échéances."),
                ("Préconditions", "L’étudiant est authentifié et a déjà saisi des données académiques exploitables."),
                ("Postconditions", "Les scores de risque et leurs facteurs explicatifs sont affichés."),
                ("Scénario nominal", "L’étudiant ouvre la page d’analyse de risque. Le système récupère les données du cours, calcule le score global et présente un niveau de criticité accompagné de détails sur la performance, la procrastination et la pression des examens."),
            ],
        ),
        (
            "CU7 : Synchroniser les données",
            [
                ("Acteurs", "API REST de synchronisation (/api/sync)"),
                ("Description", "Permet de réconcilier les données locales de la PWA avec les données persistées côté serveur après une période hors ligne."),
                ("Préconditions", "L’étudiant est authentifié, des mutations locales sont en attente et la connectivité réseau est rétablie."),
                ("Postconditions", "Les files locales sont vidées en cas de succès et l’état local redevient cohérent avec le serveur."),
                ("Scénario nominal", "À la reconnexion, le module de synchronisation pousse vers l’endpoint /api/sync les opérations locales marquées comme en attente, récupère ensuite les données à jour, met à jour le cache local et conserve uniquement les actions ayant échoué de façon temporaire."),
            ],
        ),
        (
            "CU8 : Recevoir des notifications",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet au système d’alerter l’étudiant sur les échéances proches ou les situations de risque nécessitant une réaction."),
                ("Préconditions", "L’étudiant a autorisé les notifications et au moins une échéance pertinente a été détectée."),
                ("Postconditions", "Une notification est reçue sur l’appareil de l’étudiant."),
                ("Scénario nominal", "Le système détecte qu’une échéance importante approche ou qu’un cours devient critique, puis envoie une notification permettant à l’étudiant d’ouvrir directement la tâche, l’événement ou le cours concerné."),
            ],
        ),
        (
            "CU9 : Gérer son profil",
            [
                ("Acteurs", "Étudiant"),
                ("Description", "Permet de consulter et de modifier ses informations personnelles ainsi que certaines préférences de l’application."),
                ("Préconditions", "L’étudiant est authentifié."),
                ("Postconditions", "Les nouvelles informations de profil sont sauvegardées."),
                ("Scénario nominal", "L’étudiant ouvre la page « Profil », modifie les informations souhaitées, soumet le formulaire et le système valide puis enregistre les nouvelles valeurs."),
            ],
        ),
    ]

    anchor: Paragraph | Table = student_intro
    for title, rows in student_cases:
        anchor = make_two_col_table(doc, anchor, original_style, title, rows)
        anchor = insert_paragraph_after(anchor, "", style=caption.style)

    admin_intro = insert_paragraph_after(
        anchor,
        "Les cas d’utilisation suivants relèvent de l’acteur administrateur. Ils appartiennent au modèle conceptuel du système, même s’ils ne sont pas encore couverts par l’interface du MVP.",
        style=caption.style,
    )
    anchor = admin_intro
    admin_cases = [
        (
            "CU10 : Gérer les utilisateurs",
            [
                ("Acteurs", "Administrateur"),
                ("Description", "Permet de superviser les comptes, de réinitialiser un accès ou de bloquer un utilisateur en cas d’anomalie."),
                ("Préconditions", "L’administrateur est authentifié avec les privilèges adéquats."),
                ("Postconditions", "L’état du compte utilisateur ciblé est mis à jour et l’opération est journalisée."),
                ("Scénario nominal", "L’administrateur ouvre l’espace d’administration, consulte la liste des comptes, effectue l’action de gestion souhaitée, puis le système enregistre cette intervention."),
            ],
        ),
        (
            "CU11 : Consulter les statistiques globales",
            [
                ("Acteurs", "Administrateur"),
                ("Description", "Permet de suivre les indicateurs globaux de disponibilité, d’usage et de synchronisation de la plateforme."),
                ("Préconditions", "L’administrateur est authentifié."),
                ("Postconditions", "Les indicateurs globaux sont affichés et exploitables pour la supervision."),
                ("Scénario nominal", "L’administrateur ouvre son tableau de bord, consulte le nombre d’utilisateurs, l’état des synchronisations et les indicateurs de santé applicative, puis identifie les anomalies nécessitant une intervention."),
            ],
        ),
    ]
    for title, rows in admin_cases:
        anchor = make_two_col_table(doc, anchor, original_style, title, rows)
        anchor = insert_paragraph_after(anchor, "", style=caption.style)

    # Diagrammes de séquence et d'activités.
    replace_text(
        find_paragraph(doc, "Gestion des événements", exact=True),
        "Gestion des événements",
    )
    replace_text(
        find_paragraph(doc, "Pour une récurrence hebdomadaire"),
        "La création d’un événement dans StudyFlow distingue deux cas principaux. Un événement unique correspond à une occurrence isolée, par exemple un examen fixé au mardi 20 mai à 10h00 ou un rendez-vous administratif ponctuel. Un événement récurrent, au contraire, représente une activité répétitive, comme un cours hebdomadaire chaque lundi de 8h00 à 10h00 pendant tout le semestre.",
    )
    replace_text(
        find_paragraph(doc, "Figure 6 Diagramme de séquence"),
        "Figure 6 : Diagramme de séquence — Gestion des événements",
    )
    replace_text(
        find_paragraph(doc, "Le diagramme illustre le cycle de vie d’une tâche"),
        "Le diagramme de la Figure 6 montre le traitement complet d’un événement depuis le frontend jusqu’à l’API REST. Pour un événement unique, le client envoie une requête POST /events contenant le titre, le type, la date, l’heure et le cours associé ; le backend vérifie la propriété du cours, persiste la ressource puis renvoie l’objet créé. Pour un événement récurrent, la logique de l’application génère plusieurs occurrences à partir d’une même saisie, puis les crée séquentiellement avant de rafraîchir la vue calendrier. Cette distinction explique la présence d’une boucle dans le diagramme et justifie le fait qu’une simple intention utilisateur puisse produire plusieurs écritures en base.",
    )

    replace_text(
        find_paragraph(doc, "Ce diagramme illustre une boucle d’amélioration continue"),
        "Ce diagramme illustre une boucle d’amélioration continue. En début de semaine, l’étudiant planifie ses cours et ses événements, puis décompose cette planification en tâches concrètes. Au fil des jours, il exécute ces tâches, enregistre ses notes et met à jour ses livrables. Chaque nouvelle donnée alimente ensuite l’analyse de risque. Si le risque reste faible, le plan est poursuivi ; s’il devient élevé, la boucle prévoit un retour vers la replanification afin de réorganiser les priorités, d’ajouter des séances de révision ou d’ajuster la charge de travail de la semaine suivante.",
    )
    replace_text(
        find_paragraph(doc, "Figure 12 Diagramme d’activité, planification hebdomadaire"),
        "Figure 12 : Diagramme d’activité — Planification hebdomadaire",
    )

    replace_text(
        find_paragraph(doc, "Figure 13 Diagramme d’activité, Gestion de la synchronisation"),
        "Figure 13 : Diagramme d’activité — Gestion de la synchronisation hors ligne",
    )
    sync_heading = find_paragraph(doc, "Gestion de la synchronisation offline", exact=True)
    sync_desc = insert_paragraph_after(
        sync_heading,
        "Le flux débute par une mutation initiée par l’utilisateur (création, modification ou suppression). Si la connexion est disponible, la requête est envoyée immédiatement à l’API REST. Sinon, l’action est stockée localement dans une file persistante avec un statut en attente. Lors du retour du réseau, le module de synchronisation relit cette file, pousse les opérations vers /api/sync, supprime les entrées validées et conserve uniquement les échecs temporaires pour une nouvelle tentative. Le diagramme décrit donc à la fois la continuité d’usage en mode hors ligne et la stratégie de reprise automatique à la reconnexion.",
        style=sync_heading.style,
    )

    replace_text(
        find_paragraph(doc, "Figure 14 Diagramme d’activité, Analyse de risque et recommandation"),
        "Figure 14 : Diagramme d’activité — Analyse de risque et recommandation",
    )
    risk_heading = find_paragraph(doc, "Analyse de risque et recommandation", exact=True)
    risk_desc = insert_paragraph_after(
        risk_heading,
        "L’analyse de risque s’exécute cours par cours. Le système collecte d’abord les notes, les tâches non terminées et la prochaine échéance critique (travail ou examen). Il calcule ensuite trois facteurs : la performance académique, obtenue à partir de la moyenne pondérée et inversée sur 100 ; la procrastination, estimée selon le nombre de tâches encore ouvertes ; et la pression temporelle, déduite de la proximité du prochain examen. Le score final combine ces trois dimensions avec les pondérations 40 % pour la performance, 30 % pour la charge de travail et 30 % pour l’urgence. Enfin, le niveau de risque est catégorisé selon des seuils explicites : LOW si le score est inférieur à 30, MEDIUM entre 30 et 59, HIGH entre 60 et 84 et CRITICAL à partir de 85.",
        style=risk_heading.style,
    )

    replace_text(
        find_paragraph(doc, "Figure 19 Diagramme d’activité, gestion des cours"),
        "Figure 19 : Diagramme d’activité — Gestion des cours",
    )
    courses_heading = find_paragraph(doc, "2.3.5.8 Gestion des cours", exact=True)
    insert_paragraph_after(
        courses_heading,
        "Ce diagramme couvre le même schéma général pour l’ajout, la modification et l’archivage d’un cours. L’étudiant ouvre la rubrique « Cours », déclenche l’action voulue, complète ou corrige les informations du formulaire, puis le système valide les données, persiste le changement et recharge la liste. L’archivage suit la même logique, à ceci près qu’il marque le cours comme inactif au lieu de le supprimer brutalement. Cette parenté de comportement justifie que les trois opérations soient décrites comme des variantes d’un même cas d’utilisation principal.",
        style=courses_heading.style,
    )

    replace_text(
        find_paragraph(doc, "Figure 20 Diagramme d’activité, Authentification"),
        "Figure 20 : Diagramme d’activité — Authentification",
    )
    auth_heading = find_paragraph(doc, "2.3.5.9 Authentification", exact=True)
    auth_desc = find_paragraph(doc, ".", exact=True)
    replace_text(
        auth_desc,
        "Le processus d’authentification commence par la saisie de l’adresse électronique et du mot de passe. Le système valide le format du formulaire, vérifie ensuite les identifiants puis ouvre la session si la réponse est positive. En cas d’échec, il affiche immédiatement un message d’erreur avant de renvoyer l’utilisateur vers la saisie des identifiants corrigés. La boucle du diagramme doit donc être comprise comme un retour après notification d’erreur, et non comme une répétition aveugle du processus complet.",
    )

    # Diagramme de classes: aligner la description avec le schéma de données.
    replace_text(
        find_paragraph(doc, "Figure 21 Diagramme de classes"),
        "Figure 21 : Diagramme de classes",
    )
    replace_text(
        find_paragraph(doc, "User : représente l’étudiant"),
        "User : représente l’utilisateur authentifié de la plateforme. Il possède ses cours, ses événements, ses tâches, ses travaux, ses notes et son historique de synchronisation.",
    )
    replace_text(
        find_paragraph(doc, "Course : unité d’enseignement"),
        "Course : unité d’enseignement servant d’ancrage aux autres objets métier. Un cours regroupe des événements, des tâches, des travaux, des notes ainsi que des types de travaux configurables.",
    )
    replace_text(
        find_paragraph(doc, "Task :  tâche de travail académique"),
        "Task : tâche de travail académique ou personnel, caractérisée par un statut (PENDING, IN_PROGRESS, COMPLETED, CANCELED), une priorité, une échéance éventuelle et, si nécessaire, un lien vers un cours ou un événement.",
    )
    replace_text(
        find_paragraph(doc, "Work : livrable académique"),
        "Work : livrable académique planifié ou soumis. Il est lié à un cours, peut être rattaché à un événement précis et porte, via workTypeLabel ou workTypeId, le type d’évaluation concerné (TP, interrogation, examen, projet, etc.).",
    )
    replace_text(
        find_paragraph(doc, "Grade : note associée à un cours"),
        "Grade : note enregistrée pour un cours et, de manière optionnelle, pour un travail précis. Dans le modèle de données réel, la note conserve à la fois son rattachement principal au cours et un lien facultatif vers Work, ce qui permet de relier directement une évaluation au livrable qui l’a produite.",
    )
    replace_text(
        find_paragraph(doc, "RiskAnalysis :"),
        "CourseWorkType : configuration des types de travaux propres à un cours (EXAMEN, INTERRO, PROJET, TD, TP, EXERCICES) avec leur pondération éventuelle.",
    )
    replace_text(
        find_paragraph(doc, "SyncAction :"),
        "SyncHistory : historique technique des opérations de synchronisation, utile pour suivre les échanges entre le cache local hors ligne et le backend.",
    )

    replace_text(
        find_paragraph(doc, "La modélisation UML du système a ensuite permis de formaliser ces besoins"),
        "La modélisation UML du système a ensuite permis de formaliser ces besoins à travers plusieurs niveaux de représentation complémentaires. Le diagramme des cas d’utilisation met en évidence les neuf cas d’utilisation centrés sur l’étudiant, complétés par deux cas de support attribués à l’administrateur. Les diagrammes de séquences décrivent le déroulement temporel des interactions critiques, notamment l’authentification, la gestion des événements, des tâches et la synchronisation hors ligne. Les diagrammes d’activités modélisent les flux de travail métier, incluant le cycle de vie des tâches et des travaux ainsi que le processus de calcul du risque académique. Enfin, le diagramme de classes présente la structure statique du domaine en restant cohérent avec les entités réellement persistées dans la base de données.",
    )

    # Interfaces utilisateur et légendes.
    replace_text(
        find_paragraph(doc, "Figure 25 : cours"),
        "Figure 25 : Interfaces de gestion des cours",
    )
    replace_text(
        find_paragraph(doc, "Cette section présente l’interface permettant à l’étudiant de gérer l’ensemble de ses unités d’enseignement"),
        "Cette section présente les interfaces permettant à l’étudiant de gérer l’ensemble de ses unités d’enseignement. Depuis le tableau de bord ou la barre latérale, l’utilisateur ouvre la rubrique « Cours », consulte la liste synthétique de ses matières, puis accède au détail d’un cours en cliquant sur une carte.",
    )
    course_after = insert_paragraph_after(
        find_paragraph(doc, "Figure 25 : Interfaces de gestion des cours"),
        "Les captures montrent d’abord la vue d’ensemble des cours, où chaque carte résume la moyenne, le nombre de notes, de tâches et d’événements associés. La dernière capture illustre la fiche détaillée d’un cours, structurée en onglets (« Notes », « Travaux », « Tâches », « Événements », « Risque »). L’ajout d’un nouveau cours se fait à partir de la carte « Ajouter un cours » visible dans cette même rubrique ; le formulaire d’édition est déclenché par l’action « Modifier » dans la fiche détaillée.",
        style=find_paragraph(doc, "Figure 25 : Interfaces de gestion des cours").style,
    )

    replace_text(
        find_paragraph(doc, "Figure 27 : Ecran planning"),
        "Figure 27 : Interfaces du planning académique",
    )
    replace_text(
        find_paragraph(doc, "L’écran décrit ici offre une visualisation temporelle"),
        "L’interface de planning est accessible depuis la barre latérale ou à partir du tableau de bord lorsque l’étudiant souhaite détailler ses événements du jour. Elle offre une visualisation temporelle complète des événements académiques associés aux cours et propose plusieurs modes de consultation.",
    )
    replace_text(
        find_paragraph(doc, "L’écran Calendrier permet à l’étudiant de visualiser"),
        "Les captures montrent successivement la vue mensuelle, la vue hebdomadaire et le formulaire d’ajout d’un événement. L’étudiant peut cliquer sur le bouton « + Événement » pour ouvrir cette fenêtre, préciser le titre, le type, le cours lié, l’horaire, le lieu et, si nécessaire, activer la récurrence hebdomadaire. La modification ou la suppression d’un événement s’effectue ensuite en sélectionnant l’occurrence concernée dans le planning, ce qui ouvre la fenêtre d’édition correspondante.",
    )

    replace_text(
        find_paragraph(doc, "Figure 28 : Ecran Taches"),
        "Figure 28 : Interfaces de gestion des tâches",
    )
    replace_text(
        find_paragraph(doc, "Cette sous-section détaille l’interface centrale de planification du travail personnel"),
        "Cette sous-section détaille l’interface centrale de planification du travail personnel. On y accède depuis la barre latérale, depuis le tableau de bord via « Voir tout » ou encore par la capture rapide affichée sur la page d’accueil.",
    )
    replace_text(
        find_paragraph(doc, "L’écran « Tâches » est le centre de planification du travail personnel"),
        "L’écran « Tâches » implémente le cas d’utilisation UC4 : Gérer les tâches. Les captures montrent le tableau Kanban principal, la fenêtre d’édition rapide d’une tâche et la capture rapide en langage naturel. La classification automatique s’appuie sur le statut de la tâche : une tâche créée manuellement ou par défaut apparaît dans « À faire » (PENDING), puis peut passer à « En cours », « Terminées » ou « Annulées ». Lorsque l’utilisateur saisit une consigne naturelle, le moteur de capture rapide tente également d’inférer la priorité, la date et parfois le cours concerné ; les filtres affichés en haut permettent ensuite de croiser cette organisation avec l’urgence, l’échéance du jour ou un cours donné. Le libellé « Général » désigne une tâche non encore rattachée à un cours ou à un événement précis.",
    )

    replace_text(
        find_paragraph(doc, "Figure 29 : Analyse de risque"),
        "Figure 29 : Interfaces d’analyse de risque",
    )
    replace_text(
        find_paragraph(doc, "L’interface d’analyse de risque académique synthétise"),
        "L’interface d’analyse de risque académique synthétise, pour chaque cours, le niveau de criticité calculé à partir des notes, des échéances et de la charge de travail. L’étudiant y accède depuis le menu latéral ou en ouvrant le détail d’un cours, via l’onglet « Risque ».",
    )
    replace_text(
        find_paragraph(doc, "L’écran « Analyse de risque » implémente le cas d’utilisation"),
        "L’écran « Analyse de risque » implémente le cas d’utilisation UC6 : Consulter le risque académique. Le score global est directement proportionnel au risque : plus il augmente, plus la situation du cours se rapproche des seuils HIGH puis CRITICAL. Concrètement, le backend calcule ce score à partir de trois composantes : la performance académique (40 %), la charge de travail restante ou procrastination (30 %) et la pression des échéances, notamment la proximité d’un examen (30 %). Les indicateurs par cours affichés dans la liste permettent donc de comprendre à la fois le niveau global et les facteurs qui l’expliquent.",
    )

    # Performance / sources.
    replace_text(
        find_paragraph(doc, "Lighthouse est un outil d’audit automatisé de Google"),
        "Lighthouse est un outil d’audit automatisé maintenu par Google qui évalue les performances, l’accessibilité et les bonnes pratiques d’une page web. Pour simuler les conditions réelles des étudiants de Goma (connexion instable, smartphone d’entrée de gamme), les tests ont été exécutés en mode mobile avec un ralentissement processeur (4x) et une émulation réseau 3G (latence 150 ms, débit descendant 1,6 Mbit/s), après un chargement initial à froid. Les seuils de référence et les principes de calcul retenus dans cette sous-section proviennent de la documentation officielle de Google Lighthouse et des recommandations Web Vitals.",
    )

    # Limites méthodologiques.
    replace_text(
        find_paragraph(doc, "L’évaluation de l’application n’a pas inclus de test utilisateur réel"),
        "L’évaluation de l’application n’a pas inclus de test utilisateur réel auprès d’étudiants de Goma, principalement parce que le prototype n’a été stabilisé qu’en fin de développement et n’a pas encore été déployé dans un cadre permettant de recruter, d’observer et de comparer un petit panel d’utilisateurs dans des conditions homogènes. À ce stade, la validation menée relève donc surtout de tests techniques et fonctionnels. Nous ne disposons pas encore de données empiriques sur l’acceptabilité, la satisfaction ou l’efficacité réelle de l’outil en usage prolongé, ce qui constitue une limite méthodologique importante et un chantier prioritaire pour la suite du travail.",
    )

    # Annexes: titres plus explicites.
    replace_text(find_paragraph(doc, "Schema de la base des données"), "Annexe A — Schéma de la base de données")
    replace_text(find_paragraph(doc, "Figure 30 : Appercu du schema prisma de studyFlow"), "Figure 30 : Aperçu du schéma Prisma de StudyFlow")
    replace_text(find_paragraph(doc, "Rapport Lighthouse"), "Annexe B — Rapport Lighthouse")
    replace_text(find_paragraph(doc, "Figure 31 : Aperçu du rapport de performance de StudyFlow avec lighthouse"), "Figure 31 : Aperçu du rapport de performance de StudyFlow avec Lighthouse")
    replace_text(find_paragraph(doc, "Extrait de code du Service Worker"), "Annexe C — Extrait de code du service worker")
    replace_text(find_paragraph(doc, "Figure 32 : Aperçu du code de configuration de la PWA"), "Figure 32 : Aperçu du code de configuration de la PWA")
    hook_caption = find_paragraph(doc, "Figure 33 ; Aperçu du hook d’installation de la PWA")
    hook_caption.insert_paragraph_before("Annexe D — Hook d’installation de la PWA", style=hook_caption.style)
    replace_text(hook_caption, "Figure 33 : Aperçu du hook d’installation de la PWA")
    sync1_caption = find_paragraph(doc, "Figure 34 : Aperçu du code de synchronisation 1")
    sync1_caption.insert_paragraph_before("Annexe E — Code de synchronisation hors ligne (partie 1)", style=sync1_caption.style)
    replace_text(sync1_caption, "Figure 34 : Aperçu du code de synchronisation hors ligne (partie 1)")
    sync2_caption = find_paragraph(doc, "Figure 35  : Aperçu du code de synchronisation 2")
    sync2_caption.insert_paragraph_before("Annexe F — Code de synchronisation hors ligne (partie 2)", style=sync2_caption.style)
    replace_text(sync2_caption, "Figure 35 : Aperçu du code de synchronisation hors ligne (partie 2)")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
